import datetime
import subprocess
import os
import time

from flask import Flask, render_template, request, send_file, redirect, flash
from flask_sqlalchemy import SQLAlchemy
from flask_login import (
    LoginManager, UserMixin,
    login_user, logout_user,
    login_required, current_user
)
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document
from openai import OpenAI

from ai.prompt_engine import load_prompt
from services.validator import validate_sections


# ================= APP SETUP =================
app = Flask(__name__)
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///database.db"
app.config["SECRET_KEY"] = os.getenv("SECRET_KEY", "dev-secret")


db = SQLAlchemy(app)

login_manager = LoginManager(app)
login_manager.login_view = "login"
# =============================================


# ================= DATABASE MODELS =================
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    email = db.Column(db.String(150), unique=True, nullable=False)
    password = db.Column(db.String(150), nullable=False)

    reports = db.relationship("Report", backref="user", lazy=True)


class Report(db.Model):
    id = db.Column(db.Integer, primary_key=True)

    title = db.Column(db.String(200))
    topic = db.Column(db.String(200))
    filename = db.Column(db.String(200))
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"))

    model_used = db.Column(db.String(50))
    prompt_version = db.Column(db.String(50))
    created_at = db.Column(db.DateTime, default=datetime.datetime.utcnow)

    quality_score = db.Column(db.Float)
    validation_status = db.Column(db.String(20))

    # PHASE 4
    version = db.Column(db.Integer, default=1)
    parent_report_id = db.Column(db.Integer, nullable=True)

    # PHASE 5
    generation_time_ms = db.Column(db.Integer)
    ai_confidence = db.Column(db.Float)
# ==================================================


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


# ================= SECTION PARSER =================
def parse_sections(text):
    sections = {
        "introduction": "",
        "problem_statement": "",
        "solution": "",
        "conclusion": ""
    }

    current = None
    for line in text.split("\n"):
        l = line.lower()
        if "introduction" in l:
            current = "introduction"
        elif "problem" in l:
            current = "problem_statement"
        elif "solution" in l:
            current = "solution"
        elif "conclusion" in l:
            current = "conclusion"
        elif current:
            sections[current] += line + " "

    return sections
# ================================================


# ================= AI CONFIG =================
client = OpenAI(
    api_key=os.getenv("OPENAI_API_KEY"),
    base_url="https://openrouter.ai/api/v1"
)

# =============================================


# ================= ROUTES =================
@app.route("/")
def root():
    return redirect("/login")


@app.route("/generate", methods=["GET", "POST"])
@login_required
def generate():
    if request.method == "POST":
        title = request.form["title"]
        topic = request.form["topic"]
        format_type = request.form["format"]

        if not topic.strip():
            flash("Report topic cannot be empty ❌", "error")
            return redirect("/generate")

        # 🔥 PHASE 4: VERSIONING
        existing_report = (
            Report.query
            .filter_by(user_id=current_user.id, topic=topic)
            .order_by(Report.version.desc())
            .first()
        )

        if existing_report:
            version = existing_report.version + 1
            parent_id = existing_report.parent_report_id or existing_report.id
        else:
            version = 1
            parent_id = None

        BASE_DIR = app.root_path
        template_map = {
            "college": os.path.join(BASE_DIR, "templates/formats/college_template.docx"),
            "ieee": os.path.join(BASE_DIR, "templates/formats/ieee_template.docx"),
            "simple": os.path.join(BASE_DIR, "templates/formats/simple_template.docx"),
        }

        template_path = template_map.get(format_type)
        if not template_path:
            flash("Template not found ❌", "error")
            return redirect("/generate")

        prompt, prompt_version = load_prompt(format_type, topic)

        try:
            # ========= MODEL SELECTION BY GENERATION TYPE =========
            model_map = {
                "college": "mistralai/mistral-7b-instruct",
                "ieee": "meta-llama/llama-3.1-8b-instruct",
                "simple": "gpt-4o-mini"
            }
            model_name = model_map.get(format_type, "mistralai/mistral-7b-instruct")
            # ======================================================

            # ⏱ PHASE 5: GENERATION TIME
            start_time = time.time()

            response = client.chat.completions.create(
                model=model_name,
                messages=[{"role": "user", "content": prompt}],
            )

            end_time = time.time()
            generation_time_ms = int((end_time - start_time) * 1000)

            content = response.choices[0].message.content
            sections = parse_sections(content)

            # ✅ GET VALIDATION RESULT FIRST
            quality_score, validation_status, _ = validate_sections(sections)

            # ✅ CAP QUALITY TO 100
            quality_score = min(100, quality_score)

            # ✅ CALCULATE AI CONFIDENCE
            ai_confidence = round(quality_score / 100, 2)

        except Exception as e:
            print(e)
            flash("AI service unavailable ❌", "error")
            return redirect("/generate")

        doc = Document(template_path)

        meta = doc.add_paragraph()
        meta.add_run(
            f"Generated By: Smart AI Report Generator\n"
            f"Author: {current_user.email}\n"
            f"Generated On: {datetime.datetime.now().strftime('%d %B %Y')}\n\n"
        ).bold = True

        for para in doc.paragraphs:
            para.text = para.text.replace("{{title}}", title)
            para.text = para.text.replace("{{content}}", content)

        # 📄 PHASE 5: EMBED AI METADATA
        doc.add_page_break()
        audit = doc.add_paragraph()
        audit.add_run("AI GENERATION METADATA\n").bold = True
        audit.add_run(
            f"\nModel Used: {model_name}"
            f"\nPrompt Version: {prompt_version}"
            f"\nQuality Score: {quality_score}"
            f"\nValidation Status: {validation_status}"
            f"\nAI Confidence: {ai_confidence}"
            f"\nGeneration Time: {generation_time_ms} ms"
            f"\nGenerated On: {datetime.datetime.now().strftime('%d %B %Y')}"
        )
        UPLOAD_DIR = os.path.join(app.instance_path, "reports")
        os.makedirs(UPLOAD_DIR, exist_ok=True)

        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        docx_path = os.path.join(
        UPLOAD_DIR,
        f"generated_report_{current_user.id}_{timestamp}.docx"
        )

        doc.save(docx_path)

        pdf_path = docx_path


        report = Report(
            title=title,
            topic=topic,
            filename=pdf_path,
            user_id=current_user.id,
            model_used=model_name,
            prompt_version=prompt_version,
            quality_score=quality_score,
            validation_status=validation_status,
            version=version,
            parent_report_id=parent_id,
            generation_time_ms=generation_time_ms,
            ai_confidence=ai_confidence
        )

        db.session.add(report)
        db.session.commit()

        flash("Report generated with full AI auditability ✅", "success")
        return redirect("/dashboard")

    return render_template("index.html")


@app.route("/dashboard")
@login_required
def dashboard():
    reports = current_user.reports
    total_reports = len(reports)

    avg_quality = round(
        sum(r.quality_score for r in reports) / total_reports, 2
    ) if reports else 0

    avg_generation_time = round(
        sum(r.generation_time_ms for r in reports) / total_reports, 2
    ) if reports else 0

    prompt_versions = [r.prompt_version for r in reports]
    most_used_prompt = (
        max(set(prompt_versions), key=prompt_versions.count)
        if reports else "N/A"
    )

    return render_template(
        "dashboard.html",
        reports=reports,
        total_reports=total_reports,
        avg_quality=avg_quality,
        avg_generation_time=avg_generation_time,
        most_used_prompt=most_used_prompt
    )


@app.route("/profile")
@login_required
def profile():
    reports = current_user.reports
    total_reports = len(reports)

    last_report_date = (
        max(r.created_at for r in reports).strftime("%d %b %Y")
        if reports else "N/A"
    )

    avg_quality = round(
        sum(r.quality_score for r in reports) / total_reports, 2
    ) if reports else 0
    avg_generation_time = round(
        sum(r.generation_time_ms for r in reports) / total_reports, 2
    ) if reports else 0

    prompt_versions = [r.prompt_version for r in reports]
    most_used_prompt = (
        max(set(prompt_versions), key=prompt_versions.count)
        if reports else "N/A"
    )

    topic_count = {}
    for r in reports:
        topic_count[r.topic] = topic_count.get(r.topic, 0) + 1

    most_used_topic = max(topic_count, key=topic_count.get) if topic_count else "N/A"

    return render_template(
        "profile.html",
        user=current_user,
        total_reports=total_reports,
        last_report_date=last_report_date,
        most_used_topic=most_used_topic,
        avg_quality=avg_quality,
        avg_generation_time=avg_generation_time,
        most_used_prompt=most_used_prompt
    )


@app.route("/download/<int:report_id>")
@login_required
def download(report_id):
    report = Report.query.get_or_404(report_id)
    if report.user_id != current_user.id:
        flash("Unauthorized ❌", "error")
        return redirect("/dashboard")
    return send_file(report.filename, as_attachment=True)


@app.route("/delete/<int:report_id>")
@login_required
def delete_report(report_id):
    report = Report.query.get_or_404(report_id)
    if report.user_id != current_user.id:
        flash("Unauthorized ❌", "error")
        return redirect("/dashboard")

    if os.path.exists(report.filename):
        os.remove(report.filename)

    db.session.delete(report)
    db.session.commit()
    flash("Report deleted 🗑", "success")
    return redirect("/dashboard")


# ================= AUTH =================
@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        email = request.form["email"]
        password = generate_password_hash(request.form["password"])

        if User.query.filter_by(email=email).first():
            flash("Email already exists ❌", "error")
            return redirect("/register")

        user = User(email=email, password=password)
        db.session.add(user)
        db.session.commit()

        login_user(user)
        return redirect("/generate")

    return render_template("register.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        user = User.query.filter_by(email=request.form["email"]).first()
        if user and check_password_hash(user.password, request.form["password"]):
            login_user(user)
            return redirect("/generate")

        flash("Invalid credentials ❌", "error")

    return render_template("login.html")


@app.route("/logout")
@login_required
def logout():
    logout_user()
    return redirect("/login")
# ==================================================


with app.app_context():
    db.create_all()


if __name__ == "__main__":
    app.run()

